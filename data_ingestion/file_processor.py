"""Functions for extracting and parsing resume files."""

import fitz
import pandas as pd
import json
import re
import logging
import os
from pdf2image import convert_from_path
import easyocr
import numpy as np
import docx
from docx.oxml.text.paragraph import CT_P
from docx.oxml.table import CT_Tbl
from typing import Union, Dict
from datetime import date, datetime
from data_ingestion.config import SAVE_DIR
from grok_work.groq_cilent import client
from Google_work.google_sheet import write_to_google_sheet
from data_ingestion.config import SPREADSHEET_ID
logger = logging.getLogger(__name__)


def extract_text_from_pdf(pdf_path: str) -> str:
    """Extract text from a PDF file with OCR fallback."""
    logger.info(f"Extracting text from PDF: {pdf_path}")
    text = ""

    try:
        with fitz.open(pdf_path) as doc:
            for page_num, page in enumerate(doc, start=1):
                page_text = page.get_text()
                if page_text:
                    print(f"Extracted from fitz (Page {page_num}):", repr(page_text[:500]))
                    text += page_text + "\n"

        if text.strip():
            return text
    except Exception as e:
        logger.error(f"Error extracting text from PDF with fitz: {e}")

    try:
        logger.info("No text found using pdfplumber, attempting OCR...")
        images = convert_from_path(pdf_path, dpi=200)
        logger.info(f"Converted PDF to {len(images)} image(s)")
        reader = easyocr.Reader(['en'], gpu=False)

        for i, image in enumerate(images):
            logger.info(f"Processing page {i + 1} with OCR...")
            image_np = np.array(image)
            results = reader.readtext(image_np)
            page_text = ' '.join([result[1] for result in results])
            print(f"Extracted from OCR (Page {i + 1}):", repr(page_text[:500]))

            if page_text.strip():
                text += page_text + "\n"
            else:
                logger.warning(f"No text extracted from OCR on page {i + 1}")

        if text.strip():
            logger.info("OCR extraction successful!")
        else:
            logger.warning("OCR extraction failed. No text found.")
        return text if text.strip() else "No text could be extracted from the PDF."
    except Exception as e:
        logger.error(f"Error during OCR extraction: {e}")
        raise


def extract_text_from_docx(docx_path: str, advanced_mode=False) -> Union[str, Dict]:
    """Extract text from a DOCX file with optional advanced parsing."""
    logger.info(f"Extracting text from DOCX: {docx_path}")

    try:
        doc = docx.Document(docx_path)
        full_text_parts = []

        if advanced_mode:
            result = {
                "full_text": "",
                "paragraphs": [],
                "tables": [],
                "headers": [],
                "ordered_content": ""
            }

        for para in doc.paragraphs:
            if not para.text.strip():
                continue
            full_text_parts.append(para.text)

            if advanced_mode:
                is_heading = para.style.name.startswith('Heading')
                formatted_text = []
                for run in para.runs:
                    text = run.text
                    if text.strip():
                        formatting = {"text": text, "bold": run.bold, "italic": run.italic}
                        formatted_text.append(formatting)
                para_info = {"text": para.text, "formatted_parts": formatted_text, "is_heading": is_heading}
                result["paragraphs"].append(para_info)
                if is_heading:
                    result["headers"].append(para.text)

        for table in doc.tables:
            table_data = []
            for row in table.rows:
                row_data = [cell.text for cell in row.cells]
                row_text = " | ".join(cell.text for cell in row.cells if cell.text.strip())
                if row_text:
                    full_text_parts.append(row_text)
                if advanced_mode:
                    table_data.append(row_data)
            if advanced_mode:
                result["tables"].append(table_data)

        if advanced_mode:
            def process_element(element):
                text_content = ""
                if isinstance(element, CT_P):
                    paragraph = docx.text.paragraph.Paragraph(element, None)
                    if paragraph.text.strip():
                        text_content += paragraph.text + "\n"
                elif isinstance(element, CT_Tbl):
                    table = docx.table.Table(element, None)
                    for row in table.rows:
                        row_text = []
                        for cell in row.cells:
                            for paragraph in cell.paragraphs:
                                if paragraph.text.strip():
                                    row_text.append(paragraph.text)
                        if row_text:
                            text_content += " | ".join(row_text) + "\n"
                return text_content

            all_elements_text = ""
            for element in doc.element.body:
                all_elements_text += process_element(element)

            result["ordered_content"] = all_elements_text
            result["full_text"] = "\n".join(full_text_parts)
            return result
        else:
            return "\n".join(full_text_parts)
    except Exception as e:
        logger.error(f"Error extracting text from DOCX: {e}")
        raise


def parse_resume(resume_text, file_name=None):
    """Parse resume text using Groq API with optional email metadata."""
    email_ctc = None
    experience_from_email = None
    if file_name:
        metadata_file = os.path.splitext(file_name)[0] + "_metadata.json"
        metadata_path = os.path.join(SAVE_DIR, metadata_file)
        if os.path.exists(metadata_path):
            try:
                with open(metadata_path, 'r') as f:
                    metadata = json.load(f)
                    if "ctc_from_email" in metadata:
                        email_ctc = metadata["ctc_from_email"]
                        experience_from_email = metadata["experience_from_email"]
                        print(f"📊 Using CTC from email for {file_name}: {email_ctc}")
            except Exception as e:
                print(f"⚠️ Warning: Could not read metadata: {e}")

    current_date = date.today().strftime("%Y-%m-%d")
    prompt = f"""
    Extract the following details from the resume.
    Provide ONLY a JSON object with EXACTLY these keys (no variations):
    
    - Name
    - Email Id
    - Contact No
    - Current Location
    - Total Experience
    - Designation
    - Skills
    - CTC info
    - No of companies worked with till today
    - Last company worked with
    - Loyalty %
    - Category  # New field added
    
    Instructions for specific fields:
    - Total Experience: Express in years.
        * If explicitly mentioned, use that value.
        * If not mentioned, calculate by summing professional position durations.
        * If a position has "Present" or no end date, use {current_date} for calculation.
    - Skills: Provide as comma-separated values.
    - Loyalty %: Calculate using formula: ((Total Years Worked / Number of Companies) / Total Career Duration) * 100
        * If career duration is zero (single company career), set to 100.
    - Category: Categorize into one of the following based on skills & designation:
        * "QA"
        * "Backend"
        * "Frontend"
        * "App Developer"
        * "DevOps"
        * "Data Engineer"
        * "ML/AI Engineer"
        * "Full Stack"
        * "Cybersecurity"
        * "Cloud Engineer"
        * "Embedded Developer"
        * "Game Developer"
        * "Blockchain Developer"
        * "Database Administrator"
        
        If unclear, infer from job roles and skills. If multiple roles exist, choose the **primary** role.
    """

    if email_ctc:
        prompt += f"""
        - CTC info: Use this exact value found in the email: "{email_ctc}" 
            * Do not try to extract CTC from the resume text.
        """
    if experience_from_email:
        prompt += f"""
        - Total Experience: Use this exact value found in the email: "{experience_from_email} years" 
            * Only calculate from resume if this value seems incorrect or inconsistent with the resume content.
        """
    prompt += f"""
        Provide the response in valid JSON format only, as a flat structure (no nested objects).
        Ensure you format the response as a proper JSON object without any explanations or text outside the JSON structure.
        
        Resume:
        {resume_text}
        """

    completion = client.make_request(
        model="gemma2-9b-it",
        messages=[{"role": "user", "content": prompt}],
        temperature=0,
        max_completion_tokens=6790,
        top_p=0.95,
        stream=True,
    )

    content = ""
    for chunk in completion:
        content += chunk.choices[0].delta.content or ""

    json_match = re.search(r'({.*})', content, re.DOTALL)
    if json_match:
        try:
            parsed_json = json.loads(json_match.group(1))
            flattened_json = {}
            for key, value in parsed_json.items():
                if isinstance(value, dict):
                    for sub_key, sub_value in value.items():
                        flattened_json[sub_key] = sub_value
                else:
                    flattened_json[key] = value
            return flattened_json
        except json.JSONDecodeError:
            return {"error": "Failed to parse JSON", "raw_response": content}
    return {"error": "No JSON found", "raw_response": content}


def process_single_resume(file_name, file_link=None, email_date=None):
    """Process a single resume file and update Google Sheets."""
    file_path = os.path.join(SAVE_DIR, file_name)
    print(f"Processing resume: {file_name}")

    try:
        if file_name.lower().endswith(".pdf"):
            resume_text = extract_text_from_pdf(file_path)
        else:
            resume_text = extract_text_from_docx(file_path, advanced_mode=False)

        max_chars = 20000
        if len(resume_text) > max_chars:
            logger.warning(f"Resume is very long ({len(resume_text)} chars). Truncating to {max_chars} chars.")
            resume_text = resume_text[:max_chars]

        parsed_data = parse_resume(resume_text, file_name=file_name)

        if "error" in parsed_data:
            print(f"❌ Error parsing {file_name}: {parsed_data['error']}")
        else:
            parsed_data["File Name"] = file_link if file_link else file_name

            if email_date:
                parsed_data["Date"] = email_date
                print(f"📅 Using exact email date for spreadsheet: {email_date}")
            else:
                metadata_file = os.path.splitext(file_name)[0] + "_metadata.json"
                metadata_path = os.path.join(SAVE_DIR, metadata_file)
                if os.path.exists(metadata_path):
                    try:
                        with open(metadata_path, 'r') as f:
                            metadata = json.load(f)
                            if "email_date" in metadata:
                                parsed_data["Date"] = metadata["email_date"]
                                print(f"📅 Using date from metadata: {metadata['email_date']}")
                    except Exception as e:
                        print(f"⚠️ Warning: Could not read date from metadata: {e}")

                if "Date" not in parsed_data:
                    parsed_data["Date"] = datetime.now().strftime("%d/%m/%Y")
                    print(f"⚠️ No date found in metadata, using today: {parsed_data['Date']}")

            if "CTC info" in parsed_data:
                print(f"💰 CTC information being added to spreadsheet: {parsed_data['CTC info']}")

            write_to_google_sheet(parsed_data, SPREADSHEET_ID)

    except Exception as e:
        print(f"❌ Error processing {file_name}: {e}")
        import traceback
        print(traceback.format_exc())