"""Functions for extracting and parsing resume files from a folder with Google Drive upload."""

import fitz
import json
import re
import logging
import os
import numpy as np
import docx
from typing import Union, Dict
from datetime import date
from pdf2image import convert_from_path
import easyocr
from groq_cilent import client
from sheet_agent import get_google_sheets_client , write_to_google_sheet
from drive_agent import upload_to_google_drive, check_folder_exists
from googleapiclient.discovery import build
from googleapiclient.http import MediaFileUpload, MediaInMemoryUpload

# Configure your settings here
SAVE_DIR = "resumes"  # Update this to your resume folder path
SPREADSHEET_ID = "19FL_4w0xQpPlOmT7CsPdTmlDl_zgDeXs_5666TToAVE"  # Your Google Sheet ID
DRIVE_FOLDER_ID = "1XdmlQ-v_h7rV319pfLGVdvcLFmNrIbeH" 

logger = logging.getLogger(__name__)

def extract_text_from_pdf(pdf_path: str) -> str:
    """Extract text from a PDF file with OCR fallback."""
    logger.info(f"Extracting text from PDF: {pdf_path}")
    text = ""

    try:
        with fitz.open(pdf_path) as doc:
            for page_num, page in enumerate(doc, start=1):
                page_text = page.get_text()
                if page_text:
                    print(f"Extracted from fitz (Page {page_num}):", repr(page_text[:500]))
                    text += page_text + "\n"

        if text.strip():
            return text
    except Exception as e:
        logger.error(f"Error extracting text from PDF with fitz: {e}")

    try:
        logger.info("No text found using fitz, attempting OCR...")
        images = convert_from_path(pdf_path, dpi=200)
        logger.info(f"Converted PDF to {len(images)} image(s)")
        reader = easyocr.Reader(['en'], gpu=False)

        for i, image in enumerate(images):
            logger.info(f"Processing page {i + 1} with OCR...")
            image_np = np.array(image)
            results = reader.readtext(image_np)
            page_text = ' '.join([result[1] for result in results])
            print(f"Extracted from OCR (Page {i + 1}):", repr(page_text[:500]))

            if page_text.strip():
                text += page_text + "\n"
            else:
                logger.warning(f"No text extracted from OCR on page {i + 1}")

        if text.strip():
            logger.info("OCR extraction successful!")
        else:
            logger.warning("OCR extraction failed. No text found.")
        return text if text.strip() else "No text could be extracted from the PDF."
    except Exception as e:
        logger.error(f"Error during OCR extraction: {e}")
        raise


def extract_text_from_docx(docx_path: str, advanced_mode=False) -> Union[str, Dict]:
    """Extract text from a DOCX file with optional advanced parsing."""
    logger.info(f"Extracting text from DOCX: {docx_path}")

    try:
        doc = docx.Document(docx_path)
        full_text_parts = []

        if advanced_mode:
            result = {
                "full_text": "",
                "paragraphs": [],
                "tables": [],
                "headers": [],
                "ordered_content": ""
            }

        for para in doc.paragraphs:
            if not para.text.strip():
                continue
            full_text_parts.append(para.text)

            if advanced_mode:
                is_heading = para.style.name.startswith('Heading')
                formatted_text = []
                for run in para.runs:
                    text = run.text
                    if text.strip():
                        formatting = {"text": text, "bold": run.bold, "italic": run.italic}
                        formatted_text.append(formatting)
                para_info = {"text": para.text, "formatted_parts": formatted_text, "is_heading": is_heading}
                result["paragraphs"].append(para_info)
                if is_heading:
                    result["headers"].append(para.text)

        for table in doc.tables:
            table_data = []
            for row in table.rows:
                row_data = [cell.text for cell in row.cells]
                row_text = " | ".join(cell.text for cell in row.cells if cell.text.strip())
                if row_text:
                    full_text_parts.append(row_text)
                if advanced_mode:
                    table_data.append(row_data)
            if advanced_mode:
                result["tables"].append(table_data)

        if advanced_mode:
            result["full_text"] = "\n".join(full_text_parts)
            return result
        else:
            return "\n".join(full_text_parts)
    except Exception as e:
        logger.error(f"Error extracting text from DOCX: {e}")
        raise


def parse_resume(resume_text):
    """Parse resume text using Groq API without email metadata."""
    current_date = date.today().strftime("%Y-%m-%d")
    prompt = f"""
    Extract the following details from the resume.
    Provide ONLY a JSON object with EXACTLY these keys (no variations):
    
    - Name
    - Email Id
    - Contact No
    - Current Location
    - Total Experience
    - Designation
    - Skills
    - CTC info
    - No of companies worked with till today
    - Last company worked with
    - Loyalty %
    - Category
    
    Instructions for specific fields:
    - Total Experience: Express in years.
        * If explicitly mentioned, use that value.
        * If not mentioned, calculate by summing professional position durations.
        * If a position has "Present" or no end date, use {current_date} for calculation.
    - Skills: Provide as comma-separated values.
    - Loyalty %: Calculate using formula: ((Total Years Worked / Number of Companies) / Total Career Duration) * 100
        * If career duration is zero (single company career), set to 100.
    - Category: Categorize into one of the following based on skills & designation:
        * "QA"
        * "Backend"
        * "Frontend"
        * "App Developer"
        * "DevOps"
        * "Data Engineer"
        * "ML/AI Engineer"
        * "Full Stack"
        * "Cybersecurity"
        * "Cloud Engineer"
        * "Embedded Developer"
        * "Game Developer"
        * "Blockchain Developer"
        * "Database Administrator"
        
        If unclear, infer from job roles and skills. If multiple roles exist, choose the **primary** role.
    
    Provide the response in valid JSON format only, as a flat structure (no nested objects).
    Ensure you format the response as a proper JSON object without any explanations or text outside the JSON structure.
    
    Resume:
    {resume_text}
    """

    completion = client.make_request(
        model="gemma2-9b-it",
        messages=[{"role": "user", "content": prompt}],
        temperature=0,
        max_completion_tokens=6790,
        top_p=0.95,
        stream=True,
    )

    content = ""
    for chunk in completion:
        content += chunk.choices[0].delta.content or ""

    json_match = re.search(r'({.*})', content, re.DOTALL)
    if json_match:
        try:
            parsed_json = json.loads(json_match.group(1))
            return parsed_json
        except json.JSONDecodeError:
            return {"error": "Failed to parse JSON", "raw_response": content}
    return {"error": "No JSON found", "raw_response": content}

def save_metadata(file_name, metadata):
    """Save metadata for a resume file."""
    metadata_file = os.path.splitext(file_name)[0] + "_metadata.json"
    metadata_path = os.path.join(SAVE_DIR, metadata_file)
    
    try:
        with open(metadata_path, 'w') as f:
            json.dump(metadata, f, indent=2)
        print(f"📝 Saved metadata for {file_name}")
        return True
    except Exception as e:
        print(f"❌ Error saving metadata: {e}")
        return False
    
def process_resume_folder(folder_path=SAVE_DIR, output_json=None, write_to_sheets=False, upload_to_drive=False):
    """Process all resume files in a folder and optionally save to JSON or Google Sheets."""
    results = []
    gc = None
    
    if not os.path.exists(folder_path):
        print(f"❌ Folder not found: {folder_path}")
        return results
    
    # Get all files in the folder
    all_files = os.listdir(folder_path)
    resume_files = [f for f in all_files if f.lower().endswith(('.pdf', '.docx', '.doc'))]
    
    print(f"Found {len(resume_files)} resume files in {folder_path}")
    
    # Initialize Google client if needed
    if upload_to_drive or write_to_sheets:
        gc = get_google_sheets_client()
        print("🔑 Initialized Google client for Drive/Sheets access")
        
        # Verify folder exists if we're uploading to Drive
        if upload_to_drive and DRIVE_FOLDER_ID:
            if not check_folder_exists(DRIVE_FOLDER_ID, gc):
                print(f"❌ Drive folder ID {DRIVE_FOLDER_ID} does not exist or is not accessible")
                print("⚠️ Will upload to root Drive folder instead")
    
    for file_name in resume_files:
        file_path = os.path.join(folder_path, file_name)
        print(f"Processing resume: {file_name}")

        try:
            # Extract text from file
            if file_name.lower().endswith(".pdf"):
                resume_text = extract_text_from_pdf(file_path)
            elif file_name.lower().endswith((".docx", ".doc")):
                resume_text = extract_text_from_docx(file_path, advanced_mode=False)
            else:
                print(f"⚠️ Unsupported file format: {file_name}")
                continue

            max_chars = 20000
            if len(resume_text) > max_chars:
                logger.warning(f"Resume is very long ({len(resume_text)} chars). Truncating to {max_chars} chars.")
                resume_text = resume_text[:max_chars]

            # Parse resume text
            parsed_data = parse_resume(resume_text)

            if "error" in parsed_data:
                print(f"❌ Error parsing {file_name}: {parsed_data['error']}")
                continue
                
            # Record today's date
            current_date = date.today().strftime("%d/%m/%Y")
            parsed_data["File Name"] = file_name
            parsed_data["Date"] = current_date
            
            # Upload to Google Drive if requested
            file_link = None
            if upload_to_drive:
                print(f"📤 Uploading {file_name} to Google Drive folder {DRIVE_FOLDER_ID}...")
                try:
                    with open(file_path, 'rb') as f:
                        file_content = f.read()
                    file_link = upload_to_google_drive(
                        file_name=file_name,
                        folder_id=DRIVE_FOLDER_ID, 
                        file_content=file_content, 
                        gc=gc,
                        save_dir=folder_path
                    )
                    if file_link:
                        parsed_data["Drive Link"] = file_link
                        print(f"🔗 Google Drive link: {file_link}")
                        
                        # Save metadata including drive link
                        metadata = {
                            "drive_link": file_link,
                            "original_file_name": file_name,  # Store original name in metadata
                            "parsed_date": current_date,
                            "folder_id": DRIVE_FOLDER_ID
                        }
                        save_metadata(file_name, metadata)
                    else:
                        print(f"❌ Failed to upload {file_name} to Google Drive")
                except Exception as upload_error:
                    print(f"❌ Error during upload of {file_name}: {upload_error}")
                    import traceback
                    print(traceback.format_exc())
            
            print(f"✅ Successfully parsed: {file_name}")
            
            # Add file to results
            results.append(parsed_data)
            
            # Optionally write to Google Sheets
            if write_to_sheets:
                write_to_google_sheet(parsed_data, SPREADSHEET_ID)
                print(f"📊 Added to Google Sheet: {file_name}")

        except Exception as e:
            print(f"❌ Error processing {file_name}: {e}")
            import traceback
            print(traceback.format_exc())
    
    # Save results to JSON if requested
    if output_json:
        with open(output_json, 'w') as f:
            json.dump(results, f, indent=2)
        print(f"💾 Saved results to {output_json}")
    
    return results

# Direct function call instead of using command-line arguments
if __name__ == "__main__":
    # Set your desired parameters here
    folder_path = SAVE_DIR  # Or specify a different folder
    output_json = "resume_results.json"  # Set to None if you don't want JSON output
    write_to_sheets = True  # Set to False if you don't want to write to Google Sheets
    upload_to_drive = True  # Set to False if you don't want to upload to Google Drive
    
    # Process resumes with the specified parameters
    results = process_resume_folder(folder_path, output_json, write_to_sheets, upload_to_drive)
    print(f"Processed {len(results)} resumes successfully")